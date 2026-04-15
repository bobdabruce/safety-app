"""
HazardIQ - Backend Server
Phase 1: Chat + Memory + File Organizer
Phase 2: File uploads (images, PDFs, docs)
"""
import os
import re
import hashlib
import base64
import logging
from datetime import datetime
from fastapi import FastAPI, HTTPException, File, UploadFile, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel
from anthropic import Anthropic
from dotenv import load_dotenv
import json
from pathlib import Path
from typing import List, Optional

# Set up logging
LOG_FILE = Path(__file__).parent / "activity.log"
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s | %(levelname)s | %(message)s',
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("HazardIQ")

# Load environment variables
load_dotenv(Path(__file__).parent.parent / ".env")

app = FastAPI(title="HazardIQ")

# Allow requests from frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Claude client
client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

# Data directory — uses /app/data on Railway (persistent volume), local path otherwise
DATA_DIR = Path(os.getenv("DATA_DIR", str(Path(__file__).parent.parent)))
DATA_DIR.mkdir(parents=True, exist_ok=True)

MEMORY_FILE = DATA_DIR / "memory.json"
DOWNLOADS_PATH = Path.home() / "Downloads"
DESKTOP_PATH = Path.home() / "Desktop"
SCREENSHOTS_PATH = Path.home() / "Desktop"
UPLOADS_PATH = DATA_DIR / "uploads"
UPLOADS_PATH.mkdir(exist_ok=True)

# Active document context — persists the last uploaded file's text for follow-up questions
active_doc: dict = {"filename": None, "text": None}

LIBRARY_PATH = DATA_DIR / "library"
LIBRARY_PATH.mkdir(exist_ok=True)

def load_memory():
    if MEMORY_FILE.exists():
        data = json.loads(MEMORY_FILE.read_text())
        # Ensure math_progress exists
        if "math_progress" not in data:
            data["math_progress"] = {
                "level": 1,
                "correct": 0,
                "incorrect": 0,
                "streak": 0,
                "weak_areas": [],
                "sessions": 0
            }
        # Ensure cards collection exists
        if "cards" not in data:
            data["cards"] = []
        # Ensure AI library exists
        if "ai_library" not in data:
            data["ai_library"] = []
        # Ensure AI learning progress exists
        if "ai_learning" not in data:
            data["ai_learning"] = {
                "topics_explored": [],
                "skills": [],
                "ideas": [],
                "resources": []
            }
        return data
    return {"conversations": [], "facts": [], "cards": [], "ai_library": [], "ai_learning": {
        "topics_explored": [],
        "skills": [],
        "ideas": [],
        "resources": []
    }, "math_progress": {
        "level": 1,
        "correct": 0,
        "incorrect": 0,
        "streak": 0,
        "weak_areas": [],
        "sessions": 0
    }}

def save_memory(memory):
    MEMORY_FILE.write_text(json.dumps(memory, indent=2))

# System prompt for HazardIQ
SYSTEM_PROMPT = """You are HazardIQ, Bob's personal AI assistant running on his MacBook.

About Bob:
- 50 years old, Canadian, living in Cambodia
- OHS student at University of Fredericton (COHSES program)
- Crypto investor (~$10K portfolio)
- Prefers short, clear communication - no jargon walls
- Has memory challenges - be proactive with reminders
- Math skills around 6th grade level - wants to improve. When helping with math:
  * Break problems into small steps
  * Use real-world examples (money, percentages, OHS calculations)
  * Be patient and encouraging
  * Start simple, build up gradually

OHS STUDY HELP:
Bob is studying OHS at University of Fredericton. When helping with OHS questions:
- READ the question carefully - does it ask for TRUE or FALSE statements?
- Think through each option logically before answering
- Explain WHY each answer is correct or incorrect
- Use your knowledge of industrial hygiene, safety management, hazard assessment
- If it's a calculation, show step-by-step work
- If you're unsure, say so - don't guess confidently

Key OHS concepts to remember:
- Hierarchy of controls: Elimination > Substitution > Engineering > Admin > PPE
- LEV (Local Exhaust Ventilation): captures at source, needs proper velocity, minimize duct bends
- TWA calculations, exposure limits, sampling methods
- Risk assessment: likelihood × severity

MATH PRACTICE MODE:
When Bob asks for math practice, follow this structure:
- Level 1: Basic addition/subtraction (single digit, then double digit)
- Level 2: Basic multiplication/division
- Level 3: Percentages and fractions
- Level 4: Word problems with money
- Level 5: Multi-step problems, OHS calculations

Rules:
- Give ONE problem at a time
- Wait for answer before giving next
- If correct: praise briefly, give next problem
- If wrong: explain step-by-step, then give similar problem
- After 5 correct in a row at 80%+ accuracy, suggest moving up
- Track: respond with [MATH_CORRECT] or [MATH_INCORRECT:topic] so system can track
- Current progress shown below

KEY DATES (calculate from current date below):
- Graduation: Oct 2026
- WorkSafe benefits expire: Feb 2027

CRITICAL: Do the math! If current date is Feb 2026:
- Oct 2026 is 8 months away (NOT weeks)
- Feb 2027 is 12 months away (NOT weeks)

Your personality:
- Direct and helpful, no fluff
- Anticipate needs, don't just react
- Challenge ideas when they have holes
- Remember everything across conversations
- Explain your reasoning

**ACCURACY RULES - VERY IMPORTANT:**
1. Think carefully before answering - don't rush
2. For OHS questions: double-check your reasoning before responding
3. If a question asks for TRUE statements, list TRUE ones. If it asks for FALSE, list FALSE ones. READ THE QUESTION.
4. Don't contradict yourself - if you say something is TRUE, don't then call it FALSE
5. If you're not 100% sure, say so - "I think..." or "I'm not certain but..."
6. When you make a mistake and Bob corrects you, acknowledge it clearly and give the RIGHT answer
7. For multiple choice: read ALL options before answering, then explain WHY each is right/wrong

The 10 Commandments (NEVER break these):
1. Never delete anything without explicit approval
2. Never send anything without Bob reviewing first
3. Never move money
4. Never speak for Bob
5. Never share data
6. Always explain what you did
7. Always show your reasoning
8. Kill switch available
9. Manual override on everything
10. Bob owns all data

YOUR CAPABILITIES - You CAN do these things:
1. **Clean Downloads** - When Bob asks to clean/organize Downloads, respond with: [CLEAN_DOWNLOADS]
2. **Clean Desktop** - When Bob asks to clean/organize Desktop, respond with: [CLEAN_DESKTOP]
3. **Clean Screenshots** - When Bob asks to clean screenshots, respond with: [CLEAN_SCREENSHOTS]
4. **Remember facts** - When Bob tells you something important, respond with [REMEMBER: fact here]

Commands to use:
- Downloads/download folder → [CLEAN_DOWNLOADS]
- Desktop → [CLEAN_DESKTOP]
- Screenshots → [CLEAN_SCREENSHOTS]

SPORTS CARD RESEARCH MODE:
Bob is a serious card collector. When he clicks "Cards" or asks about cards, show this menu:

"🃏 **Card Research**

**Price Check:**
- Drop a card photo or type card details
- I'll search eBay sold listings for real prices

**What I can help with:**
- Price lookups (any card, any sport)
- Card identification from photos
- Market trends and advice
- Grading info (PSA, BGS, SGC)
- Buying/selling strategy

**Just ask:**
- 'price on 2020 Prizm Herbert PSA 10'
- 'what's this card worth?' + photo
- 'should I grade this card?'

What card do you want to research?"

**CRITICAL - YOU HAVE REAL PRICE LOOKUP!**
You CAN search eBay and get REAL prices. When Bob asks about ANY card value:

1. ALWAYS respond with: [PRICE_CHECK: search terms]
2. The system scrapes REAL sold listings from eBay
3. You will get back actual recent sale prices with clickable links

**SEARCH TERM RULES - BE SPECIFIC:**
- Include: year, brand, player name, card type
- For graded: add "PSA 10" or "BGS 9.5" etc.
- For parallels: add "silver", "refractor", "prizm" etc.
- For numbered cards: add "/25" or "/99" etc.
- For autos: add "auto" or "autograph"
- KEEP IT SIMPLE - don't add unnecessary words

Examples:
- "Herbert rookie" → [PRICE_CHECK: 2020 Prizm Herbert rookie]
- "Herbert PSA 10" → [PRICE_CHECK: 2020 Prizm Herbert PSA 10]
- "Jordan rookie raw" → [PRICE_CHECK: 1986 Fleer Jordan rookie]
- "Ramirez auto /25" → [PRICE_CHECK: 2022 Topps Archives Ramirez auto /25]
- "Luka silver prizm" → [PRICE_CHECK: 2018 Prizm Luka Doncic silver]

**IMPORTANT:**
- If results look wrong, try a DIFFERENT search (shorter or more specific)
- Ask Bob to clarify if the card details are unclear
- NEVER say "I can't search" - YOU CAN with [PRICE_CHECK:]
- NEVER give vague price ranges - get REAL data

AI LEARNING MODE:
When Bob asks about AI learning, clicks the AI button, or says "AI learning menu", help him with:

**Show this menu:**
"🤖 **AI Learning Hub**

**Learn:**
1. What is AI and how does it work?
2. Tools I can use today (ChatGPT, Claude, Midjourney, etc.)
3. AI for my specific situation (OHS, cards, Cambodia)

**Make Money with AI:**
4. Freelance opportunities (writing, images, automation)
5. Build AI-powered products/services
6. Content creation with AI
7. AI consulting for small businesses

**My Library:**
8. View my saved AI documents
9. Add a document to library (drag & drop)

**Track Progress:**
10. What AI skills have I learned?

What interests you? (Pick a number or ask anything)"

When Bob uploads documents about AI, save them to his library with [SAVE_TO_LIBRARY: title].
When he learns something new, track it with [AI_SKILL: skill learned].
When he has a business idea, save it with [AI_IDEA: idea description].

Be practical - focus on things Bob can actually do with his situation:
- 50 years old, not super technical
- Living in Cambodia (remote work friendly)
- Has OHS background (safety consulting?)
- Interested in sports cards (niche expertise)
- Wants supplemental income, not a full career change

You have real access to Bob's Mac. Always preview first, then ask before proceeding.

Current date: {date}
"""

class ChatRequest(BaseModel):
    message: str

class ChatResponse(BaseModel):
    response: str
    timestamp: str

class OrganizeResponse(BaseModel):
    message: str
    files_moved: list
    preview_mode: bool

class DuplicatesResponse(BaseModel):
    message: str
    duplicates: list
    total_size_mb: float
    preview_mode: bool

@app.get("/")
async def root():
    return FileResponse(Path(__file__).parent.parent / "frontend" / "index.html")

@app.get("/guide")
@app.get("/guide/")
async def guide():
    return FileResponse(Path(__file__).parent.parent / "frontend" / "guide" / "index.html")

@app.post("/api/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """Main chat endpoint"""
    logger.info(f"CHAT REQUEST: {request.message[:100]}")
    memory = load_memory()

    # CHECK FOR APPROVAL FIRST - before calling Claude
    # STRICT CHECK: Only trigger cleanup if the EXACT cleanup preview phrase is in the last message
    # This prevents accidental cleanup when user says "yes" to unrelated questions
    if request.message.lower() in ["yes", "yes, organize them", "organize them", "do it", "proceed", "go ahead", "yes organize them", "yes, delete them", "delete them", "yes delete them", "yes, clean them", "clean them"]:
        last_content = memory["conversations"][-1].get("content", "").lower() if memory["conversations"] else ""
        logger.info(f"APPROVAL DETECTED - checking last content: {last_content[:100]}")

        # CRITICAL: Only proceed if this was a cleanup preview (contains the exact confirmation prompt)
        is_cleanup_preview = "say **'yes'** to proceed" in last_content or "say **'yes, organize them'**" in last_content

        if not is_cleanup_preview:
            logger.info("NOT a cleanup preview - passing to Claude instead")
        else:
            # Detect which location was being cleaned
            target_path = None
            location_name = None
            if "download" in last_content:
                target_path = DOWNLOADS_PATH
                location_name = "Downloads"
            elif "desktop" in last_content:
                target_path = DESKTOP_PATH
                location_name = "Desktop"
            elif "screenshot" in last_content:
                target_path = DESKTOP_PATH  # Screenshots are on Desktop by default
                location_name = "Screenshots"

            if target_path:
                logger.info(f"EXECUTING CLEANUP on {location_name} - User approved cleanup preview")
                # Do both: organize files AND remove duplicates
                org_result = organize_folder(target_path, preview=False)
                dup_result = find_duplicates_in(target_path, preview=False)
                logger.info(f"CLEANUP COMPLETE: {len(org_result.files_moved)} organized, {len(dup_result.duplicates)} duplicates removed")

                parts = []
                if org_result.files_moved:
                    parts.append(f"**{len(org_result.files_moved)} files** sorted into folders")
                if dup_result.duplicates:
                    parts.append(f"**{len(dup_result.duplicates)} duplicates** moved to Trash ({dup_result.total_size_mb:.1f} MB)")

                if parts:
                    assistant_message = "Done! " + " and ".join(parts) + f".\n\nYour {location_name} is clean. Check Trash if you need to restore anything."
                else:
                    assistant_message = f"Nothing to do - your {location_name} was already clean."

                memory["conversations"].append({
                    "role": "user",
                    "content": request.message,
                    "timestamp": datetime.now().isoformat()
                })
                memory["conversations"].append({
                    "role": "assistant",
                    "content": assistant_message,
                    "timestamp": datetime.now().isoformat()
                })
                save_memory(memory)

                return ChatResponse(
                    response=assistant_message,
                    timestamp=datetime.now().isoformat()
                )

    # Build conversation history for context
    messages = []

    # Add recent conversation history (last 10 exchanges)
    for conv in memory["conversations"][-20:]:
        messages.append({"role": conv["role"], "content": conv["content"]})

    # Add current message
    messages.append({"role": "user", "content": request.message})

    # Build system prompt with current date
    system = SYSTEM_PROMPT.format(date=datetime.now().strftime("%Y-%m-%d %H:%M"))

    # Add remembered facts to system prompt
    if memory["facts"]:
        system += "\n\nThings you remember about Bob:\n"
        for fact in memory["facts"][-20:]:
            system += f"- {fact}\n"

    # Add math progress
    mp = memory.get("math_progress", {})
    if mp:
        total = mp.get("correct", 0) + mp.get("incorrect", 0)
        accuracy = (mp.get("correct", 0) / total * 100) if total > 0 else 0
        system += f"\n\nMath Progress: Level {mp.get('level', 1)}, {mp.get('correct', 0)} correct, {mp.get('incorrect', 0)} incorrect ({accuracy:.0f}% accuracy), streak: {mp.get('streak', 0)}"
        if mp.get("weak_areas"):
            system += f", weak areas: {', '.join(mp['weak_areas'][-3:])}"

    # Add AI learning progress
    ai_learning = memory.get("ai_learning", {})
    ai_library = memory.get("ai_library", [])
    if ai_learning.get("skills") or ai_library:
        system += f"\n\nAI Learning: {len(ai_learning.get('skills', []))} skills learned, {len(ai_library)} docs in library"
        if ai_learning.get("skills"):
            system += f"\nSkills: {', '.join(ai_learning['skills'][-5:])}"
        if ai_learning.get("ideas"):
            system += f"\nIdeas saved: {len(ai_learning['ideas'])}"

    # Add card collection summary
    cards = memory.get("cards", [])
    if cards:
        total_value = sum(c.get("estimated_value", 0) or 0 for c in cards)
        total_cost = sum(c.get("purchase_price", 0) or 0 for c in cards)
        system += f"\n\nCard Collection: {len(cards)} cards, est. value ${total_value:.2f}, cost ${total_cost:.2f}, P/L ${total_value - total_cost:.2f}"
        sorted_cards = sorted(cards, key=lambda x: x.get("estimated_value", 0) or 0, reverse=True)[:5]
        if sorted_cards:
            system += "\nTop cards: " + ", ".join([f"{c.get('name', 'Unknown')} (${c.get('estimated_value', 0)})" for c in sorted_cards])

    # Inject active document context so follow-up questions can reference the uploaded file
    if active_doc["text"]:
        system += f"\n\nACTIVE DOCUMENT — the user previously uploaded '{active_doc['filename']}'. Its full content is below. Answer questions about it directly without asking the user to re-upload or copy/paste.\n\n---\n{active_doc['text'][:40000]}\n---"

    try:
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1024,
            system=system,
            messages=messages
        )

        assistant_message = response.content[0].text

        # Check for special commands - location cleaning
        clean_commands = {
            "[CLEAN_DOWNLOADS]": (DOWNLOADS_PATH, "Downloads"),
            "[CLEAN_DESKTOP]": (DESKTOP_PATH, "Desktop"),
            "[CLEAN_SCREENSHOTS]": (DESKTOP_PATH, "Screenshots"),
        }

        for cmd, (target_path, location_name) in clean_commands.items():
            if cmd in assistant_message:
                # Execute immediately - no confirmation needed
                org_result = organize_folder(target_path, preview=False)
                dup_result = find_duplicates_in(target_path, preview=False)

                parts = []
                if org_result.files_moved:
                    parts.append(f"**{len(org_result.files_moved)} files** sorted into folders")
                if dup_result.duplicates:
                    parts.append(f"**{len(dup_result.duplicates)} duplicates** moved to Trash ({dup_result.total_size_mb:.1f} MB)")

                if parts:
                    assistant_message = "Done! " + " and ".join(parts) + f".\n\nYour {location_name} is clean."
                else:
                    assistant_message = f"Your {location_name} is already clean! Nothing to organize."
                break

        # Auto-detect if user asked about card price but Claude didn't use [PRICE_CHECK:]
        price_keywords = ["what's", "whats", "how much", "price", "worth", "value", "selling for", "going for"]
        card_keywords = ["card", "rookie", "prizm", "topps", "panini", "upper deck", "fleer", "psa", "bgs", "auto", "refractor"]
        user_msg_lower = request.message.lower()

        user_asked_price = any(kw in user_msg_lower for kw in price_keywords)
        user_mentioned_card = any(kw in user_msg_lower for kw in card_keywords)

        if user_asked_price and user_mentioned_card and "[PRICE_CHECK:" not in assistant_message:
            # Claude forgot to use price check - extract card details and search anyway
            logger.info(f"AUTO PRICE CHECK - Claude forgot command, user asked: {request.message}")

            # Use the user's message as search query (clean it up)
            search_query = request.message
            for word in ["what's", "whats", "how much is", "price on", "worth", "value of", "find", "search", "look up", "?", "the"]:
                search_query = search_query.lower().replace(word, "")
            search_query = " ".join(search_query.split())  # Clean whitespace

            if len(search_query) > 5:
                sales = await scrape_card_prices(search_query)

                if sales:
                    prices = sorted([s["price"] for s in sales])
                    median = prices[len(prices) // 2]

                    # Show top 5 with links
                    links = [f"[${s['price']:.2f}]({s['url']})" for s in sales[:5]]

                    price_report = f"""\n\n**{len(sales)} eBay sales found:**
💰 ${min(prices):.2f} - ${max(prices):.2f} | **Median: ${median:.2f}**
🔗 {' | '.join(links)}"""

                    assistant_message += price_report

        # Check for price check command
        if "[PRICE_CHECK:" in assistant_message:
            # Extract search query
            price_match = re.search(r'\[PRICE_CHECK:\s*([^\]]+)\]', assistant_message)
            if price_match:
                search_query = price_match.group(1).strip()
                logger.info(f"PRICE CHECK triggered for: {search_query}")

                # Run the actual scraper - now returns list of sales with URLs
                sales = await scrape_card_prices(search_query)

                if sales:
                    prices = sorted([s["price"] for s in sales])
                    low = min(prices)
                    high = max(prices)
                    median = prices[len(prices) // 2]

                    # Build listing links for verification - clean titles
                    listings = []
                    for sale in sales[:10]:
                        clean_title = sale['title'][:45].replace('[', '').replace(']', '')
                        clean_url = sale['url'].split('?')[0]  # Remove query params for cleaner URL
                        listings.append(f"• **${sale['price']:.2f}** - [{clean_title}]({clean_url})")

                    listings_text = "\n".join(listings)
                    ebay_search = f"https://www.ebay.com/sch/i.html?_nkw={search_query.replace(' ', '+')}&LH_Complete=1&LH_Sold=1"

                    # Build price report with individual links
                    price_report = f"""**eBay Sold: "{search_query}"**

💰 **{len(sales)} sales** | Low: ${low:.2f} | **Median: ${median:.2f}** | High: ${high:.2f}

**Recent sales (click to verify):**
{listings_text}

🔗 [All eBay results]({ebay_search})

*If these don't look right, tell me more details about your card and I'll search again.*"""

                    assistant_message = re.sub(r'\[PRICE_CHECK:[^\]]+\]', price_report, assistant_message)
                else:
                    no_results = f"""**No recent sales found for "{search_query}"**

Try different search terms or check manually:
🔗 [Search eBay Sold](https://www.ebay.com/sch/i.html?_nkw={search_query.replace(' ', '+')}&LH_Complete=1&LH_Sold=1)"""
                    assistant_message = re.sub(r'\[PRICE_CHECK:[^\]]+\]', no_results, assistant_message)

        # Check for AI library save command
        if "[SAVE_TO_LIBRARY:" in assistant_message:
            lib_match = re.findall(r'\[SAVE_TO_LIBRARY:\s*([^\]]+)\]', assistant_message)
            for title in lib_match:
                memory["ai_library"].append({
                    "title": title.strip(),
                    "added": datetime.now().isoformat()
                })
                logger.info(f"LIBRARY: Added '{title}'")
            assistant_message = re.sub(r'\[SAVE_TO_LIBRARY:[^\]]+\]', '', assistant_message).strip()

        # Check for AI skill tracking
        if "[AI_SKILL:" in assistant_message:
            skill_match = re.findall(r'\[AI_SKILL:\s*([^\]]+)\]', assistant_message)
            for skill in skill_match:
                if skill.strip() not in memory["ai_learning"]["skills"]:
                    memory["ai_learning"]["skills"].append(skill.strip())
                    logger.info(f"AI SKILL: Learned '{skill}'")
            assistant_message = re.sub(r'\[AI_SKILL:[^\]]+\]', '', assistant_message).strip()

        # Check for AI idea tracking
        if "[AI_IDEA:" in assistant_message:
            idea_match = re.findall(r'\[AI_IDEA:\s*([^\]]+)\]', assistant_message)
            for idea in idea_match:
                memory["ai_learning"]["ideas"].append({
                    "idea": idea.strip(),
                    "added": datetime.now().isoformat()
                })
                logger.info(f"AI IDEA: Saved '{idea}'")
            assistant_message = re.sub(r'\[AI_IDEA:[^\]]+\]', '', assistant_message).strip()

        # Check for remember command
        if "[REMEMBER:" in assistant_message:
            facts = re.findall(r'\[REMEMBER: (.*?)\]', assistant_message)
            for fact in facts:
                memory["facts"].append(fact)
            assistant_message = re.sub(r'\[REMEMBER: .*?\]', '', assistant_message).strip()

        # Track math progress
        if "[MATH_CORRECT]" in assistant_message:
            memory["math_progress"]["correct"] += 1
            memory["math_progress"]["streak"] += 1
            # Level up after 5 correct in a row
            if memory["math_progress"]["streak"] >= 5 and memory["math_progress"]["level"] < 5:
                total = memory["math_progress"]["correct"] + memory["math_progress"]["incorrect"]
                accuracy = memory["math_progress"]["correct"] / total if total > 0 else 0
                if accuracy >= 0.8:
                    memory["math_progress"]["level"] += 1
                    memory["math_progress"]["streak"] = 0
            assistant_message = re.sub(r'\[MATH_CORRECT\]', '', assistant_message).strip()

        if "[MATH_INCORRECT:" in assistant_message:
            memory["math_progress"]["incorrect"] += 1
            memory["math_progress"]["streak"] = 0
            # Track weak area
            weak = re.findall(r'\[MATH_INCORRECT:(.*?)\]', assistant_message)
            for area in weak:
                if area not in memory["math_progress"]["weak_areas"]:
                    memory["math_progress"]["weak_areas"].append(area)
            assistant_message = re.sub(r'\[MATH_INCORRECT:.*?\]', '', assistant_message).strip()

        # Save to memory
        memory["conversations"].append({
            "role": "user",
            "content": request.message,
            "timestamp": datetime.now().isoformat()
        })
        memory["conversations"].append({
            "role": "assistant",
            "content": assistant_message,
            "timestamp": datetime.now().isoformat()
        })

        # Keep only last 100 messages to manage size
        memory["conversations"] = memory["conversations"][-100:]
        save_memory(memory)

        return ChatResponse(
            response=assistant_message,
            timestamp=datetime.now().isoformat()
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/chat/upload", response_model=ChatResponse)
async def chat_with_upload(
    message: str = Form(default="What is this?"),
    files: List[UploadFile] = File(...)
):
    """Chat endpoint with file uploads (images, PDFs, etc.)"""
    memory = load_memory()

    # Process uploaded files
    content_parts = []
    file_descriptions = []

    for file in files:
        file_bytes = await file.read()
        file_ext = Path(file.filename).suffix.lower()

        # Save file to uploads folder
        save_path = UPLOADS_PATH / f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file.filename}"
        save_path.write_bytes(file_bytes)

        if file_ext in ['.jpg', '.jpeg', '.png', '.gif', '.webp', '.heic']:
            # Image - use Claude vision (compress if >4MB)
            from io import BytesIO
            try:
                from PIL import Image
                img = Image.open(BytesIO(file_bytes))

                # Convert HEIC or other formats to JPEG
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')

                # Resize if too large (max 4MB for safety margin)
                max_size = 4 * 1024 * 1024  # 4MB
                quality = 85
                output = BytesIO()

                while True:
                    output.seek(0)
                    output.truncate()
                    img.save(output, format='JPEG', quality=quality)
                    if output.tell() <= max_size or quality <= 20:
                        break
                    # Reduce quality or size
                    if quality > 40:
                        quality -= 15
                    else:
                        # Resize image
                        img = img.resize((int(img.width * 0.7), int(img.height * 0.7)), Image.LANCZOS)
                        quality = 70

                file_bytes = output.getvalue()
                media_type = "image/jpeg"
            except ImportError:
                # No PIL, try without compression
                media_type = f"image/{file_ext[1:]}"
                if file_ext == '.jpg':
                    media_type = "image/jpeg"

            base64_image = base64.standard_b64encode(file_bytes).decode('utf-8')

            content_parts.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": base64_image
                }
            })
            file_descriptions.append(f"[Image: {file.filename}]")

        elif file_ext == '.pdf':
            try:
                import fitz  # PyMuPDF
                pdf = fitz.open(stream=file_bytes, filetype="pdf")
                text = "".join(page.get_text() for page in pdf)
                pdf.close()
                active_doc["filename"] = file.filename
                active_doc["text"] = text
                content_parts.append({
                    "type": "text",
                    "text": f"[PDF Content from {file.filename}]:\n{text[:40000]}"
                })
                file_descriptions.append(f"[PDF: {file.filename}]")
            except ImportError:
                content_parts.append({
                    "type": "text",
                    "text": f"[PDF uploaded: {file.filename} - install PyMuPDF to read content]"
                })
                file_descriptions.append(f"[PDF: {file.filename}]")

        elif file_ext == '.docx':
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(file_bytes))
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                active_doc["filename"] = file.filename
                active_doc["text"] = text
                content_parts.append({
                    "type": "text",
                    "text": f"[Word Document from {file.filename}]:\n{text[:40000]}"
                })
                file_descriptions.append(f"[Word: {file.filename}]")
            except Exception as e:
                file_descriptions.append(f"[Word: {file.filename} - could not read: {e}]")

        elif file_ext in ['.txt', '.md', '.csv', '.json']:
            try:
                text_content = file_bytes.decode('utf-8')
                active_doc["filename"] = file.filename
                active_doc["text"] = text_content
                content_parts.append({
                    "type": "text",
                    "text": f"[File: {file.filename}]:\n{text_content[:40000]}"
                })
                file_descriptions.append(f"[Text: {file.filename}]")
            except:
                file_descriptions.append(f"[File: {file.filename} - could not read]")

        else:
            file_descriptions.append(f"[File: {file.filename} - unsupported format]")

    # Add the user's message
    content_parts.append({"type": "text", "text": message})

    # Build system prompt
    system = SYSTEM_PROMPT.format(date=datetime.now().strftime("%Y-%m-%d %H:%M"))
    if memory["facts"]:
        system += "\n\nThings you remember about Bob:\n"
        for fact in memory["facts"][-20:]:
            system += f"- {fact}\n"

    try:
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=2048,
            system=system,
            messages=[{"role": "user", "content": content_parts}]
        )

        assistant_message = response.content[0].text

        # Check for price check command in upload response
        if "[PRICE_CHECK:" in assistant_message:
            price_match = re.search(r'\[PRICE_CHECK:\s*([^\]]+)\]', assistant_message)
            if price_match:
                search_query = price_match.group(1).strip()
                logger.info(f"UPLOAD PRICE CHECK triggered for: {search_query}")

                sales = await scrape_card_prices(search_query)

                if sales:
                    prices = sorted([s["price"] for s in sales])
                    median = prices[len(prices) // 2]

                    listings = [f"• **${s['price']:.2f}** - [{s['title'][:40]}...]({s['url']})" for s in sales[:8]]

                    price_report = f"""**eBay Sold Listings for "{search_query}"**

💰 **{len(sales)} sales** | ${min(prices):.2f} - ${max(prices):.2f} | **Median: ${median:.2f}**

**Verify:**
{chr(10).join(listings)}

🔗 [All results](https://www.ebay.com/sch/i.html?_nkw={search_query.replace(' ', '+')}&LH_Complete=1&LH_Sold=1)"""

                    assistant_message = re.sub(r'\[PRICE_CHECK:[^\]]+\]', price_report, assistant_message)
                else:
                    no_results = f"""**No sales found for "{search_query}"**
🔗 [Search eBay](https://www.ebay.com/sch/i.html?_nkw={search_query.replace(' ', '+')}&LH_Complete=1&LH_Sold=1)"""
                    assistant_message = re.sub(r'\[PRICE_CHECK:[^\]]+\]', no_results, assistant_message)

        # Auto-detect card image + price request
        price_words = ["price", "worth", "value", "how much", "what's it", "find"]
        if any(pw in message.lower() for pw in price_words) and "[PRICE_CHECK:" not in assistant_message:
            # Claude identified card but forgot price check - try to extract card details
            card_patterns = [
                r'(\d{4}[-\s]?\d{2,4}?\s+\w+.*?(?:rookie|rc|auto|prizm|topps|panini|upper deck|fleer)[^\n]*)',
                r'((?:PSA|BGS)\s*\d+[^\n]*)',
            ]
            for pattern in card_patterns:
                match = re.search(pattern, assistant_message, re.IGNORECASE)
                if match:
                    search_query = match.group(1).strip()[:100]
                    logger.info(f"AUTO PRICE from upload: {search_query}")

                    sales = await scrape_card_prices(search_query)
                    if sales:
                        prices = sorted([s["price"] for s in sales])
                        median = prices[len(prices) // 2]
                        links = [f"[${s['price']:.2f}]({s['url']})" for s in sales[:5]]
                        assistant_message += f"\n\n**{len(sales)} sales:** {' | '.join(links)} | Median: ${median:.2f}"
                    break

        # Auto-save card identifications to memory
        # When Claude identifies a card from an uploaded image, save it as a fact
        card_id_patterns = [
            r'(\d{4}[-\s]?\d{2,4}?\s+(?:Topps|Prizm|Panini|Upper Deck|Fleer|Bowman)[^.!?\n]+(?:rookie|rc|auto|card)[^.!?\n]*)',
            r'((?:PSA|BGS)\s*\d+[^.!?\n]+)',
            r'This is a[n]?\s+(\d{4}[^.!?\n]+card)',
        ]
        for pattern in card_id_patterns:
            match = re.search(pattern, assistant_message, re.IGNORECASE)
            if match:
                card_desc = match.group(1).strip()[:150]
                # Save to facts for later recall
                today = datetime.now().strftime("%Y-%m-%d")
                fact = f"Card shown on {today}: {card_desc}"
                if fact not in memory.get("facts", []):
                    memory["facts"].append(fact)
                    logger.info(f"CARD MEMORY SAVED: {card_desc[:60]}")
                break

        # Save to memory
        user_msg = f"{message} {' '.join(file_descriptions)}"
        memory["conversations"].append({
            "role": "user",
            "content": user_msg,
            "timestamp": datetime.now().isoformat()
        })
        memory["conversations"].append({
            "role": "assistant",
            "content": assistant_message,
            "timestamp": datetime.now().isoformat()
        })
        memory["conversations"] = memory["conversations"][-100:]
        save_memory(memory)

        return ChatResponse(
            response=assistant_message,
            timestamp=datetime.now().isoformat()
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/remember")
async def remember_fact(request: ChatRequest):
    """Store a fact about Bob"""
    memory = load_memory()
    memory["facts"].append(request.message)
    memory["facts"] = memory["facts"][-50:]  # Keep last 50 facts
    save_memory(memory)
    return {"status": "remembered", "fact": request.message}

@app.get("/api/memory")
async def get_memory():
    """View all stored memory (Bob's control)"""
    return load_memory()

@app.delete("/api/memory")
async def clear_memory():
    """Clear all memory (Bob's control)"""
    save_memory({"conversations": [], "facts": [], "math_progress": {
        "level": 1, "correct": 0, "incorrect": 0, "streak": 0, "weak_areas": [], "sessions": 0
    }})
    return {"status": "memory cleared"}

@app.get("/api/math/progress")
async def get_math_progress():
    """Get math practice progress"""
    memory = load_memory()
    mp = memory.get("math_progress", {})
    total = mp.get("correct", 0) + mp.get("incorrect", 0)
    accuracy = (mp.get("correct", 0) / total * 100) if total > 0 else 0
    return {
        "level": mp.get("level", 1),
        "correct": mp.get("correct", 0),
        "incorrect": mp.get("incorrect", 0),
        "accuracy": round(accuracy, 1),
        "streak": mp.get("streak", 0),
        "weak_areas": mp.get("weak_areas", [])
    }

@app.delete("/api/math/progress")
async def reset_math_progress():
    """Reset math progress to start over"""
    memory = load_memory()
    memory["math_progress"] = {
        "level": 1, "correct": 0, "incorrect": 0, "streak": 0, "weak_areas": [], "sessions": 0
    }
    save_memory(memory)
    return {"status": "math progress reset"}

@app.get("/api/logs")
async def get_activity_logs():
    """View recent activity logs"""
    if LOG_FILE.exists():
        lines = LOG_FILE.read_text().split('\n')
        return {"logs": lines[-50:]}  # Last 50 lines
    return {"logs": []}

@app.delete("/api/logs")
async def clear_logs():
    """Clear activity logs"""
    if LOG_FILE.exists():
        LOG_FILE.write_text("")
    return {"status": "logs cleared"}

@app.get("/api/organize/preview")
async def preview_organize():
    """Preview what files would be organized (no action taken)"""
    return organize_downloads(preview=True)

@app.post("/api/organize")
async def do_organize():
    """Actually organize the Downloads folder (requires explicit call)"""
    return organize_downloads(preview=False)

def organize_folder(folder_path, preview=True):
    """Organize any folder into categories"""

    categories = {
        "Images": [".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg", ".heic", ".bmp", ".tiff"],
        "Documents": [".pdf", ".doc", ".docx", ".txt", ".rtf", ".odt", ".xls", ".xlsx", ".ppt", ".pptx", ".pages", ".numbers", ".key"],
        "Videos": [".mp4", ".mov", ".avi", ".mkv", ".wmv", ".webm", ".m4v"],
        "Audio": [".mp3", ".wav", ".aac", ".flac", ".m4a", ".ogg"],
        "Archives": [".zip", ".rar", ".7z", ".tar", ".gz", ".dmg"],
        "Code": [".py", ".js", ".html", ".css", ".json", ".xml", ".md", ".ts", ".jsx", ".tsx"],
        "Installers": [".pkg", ".app", ".exe", ".msi"],
    }

    files_to_move = []

    if not folder_path.exists():
        return OrganizeResponse(
            message="Folder not found",
            files_moved=[],
            preview_mode=preview
        )

    for file in folder_path.iterdir():
        if file.is_file() and not file.name.startswith('.'):
            ext = file.suffix.lower()
            for category, extensions in categories.items():
                if ext in extensions:
                    dest_folder = folder_path / category
                    files_to_move.append({
                        "file": file.name,
                        "from": str(file),
                        "to": str(dest_folder / file.name),
                        "category": category
                    })

                    if not preview:
                        dest_folder.mkdir(exist_ok=True)
                        logger.info(f"MOVE FILE: {file.name} → {category}/")
                        file.rename(dest_folder / file.name)
                    break

    if preview:
        msg = f"Preview: {len(files_to_move)} files would be organized."
    else:
        logger.info(f"ORGANIZE COMPLETE: {len(files_to_move)} files moved in {folder_path}")
        msg = f"Organized {len(files_to_move)} files into categories."

    return OrganizeResponse(
        message=msg,
        files_moved=files_to_move,
        preview_mode=preview
    )

def get_file_hash(filepath, chunk_size=8192):
    """Get MD5 hash of file for comparison"""
    hasher = hashlib.md5()
    try:
        with open(filepath, 'rb') as f:
            while chunk := f.read(chunk_size):
                hasher.update(chunk)
        return hasher.hexdigest()
    except:
        return None

def move_to_trash(filepath):
    """Move file to macOS Trash (recoverable)"""
    trash_path = Path.home() / ".Trash"
    dest = trash_path / filepath.name
    # Handle name conflicts in Trash
    counter = 1
    while dest.exists():
        dest = trash_path / f"{filepath.stem}_{counter}{filepath.suffix}"
        counter += 1
    filepath.rename(dest)
    return dest

def find_duplicates_in(folder_path, preview=True):
    """Find duplicate files in any folder"""
    if not folder_path.exists():
        return DuplicatesResponse(
            message="Folder not found",
            duplicates=[],
            total_size_mb=0,
            preview_mode=preview
        )

    # Build hash map of all files (skip tiny files under 1KB)
    MIN_SIZE = 1024  # 1KB minimum
    hash_map = {}
    for item in folder_path.rglob("*"):
        if item.is_file() and not item.name.startswith('.'):
            try:
                if item.stat().st_size < MIN_SIZE:
                    continue  # Skip tiny/empty files
            except:
                continue
            file_hash = get_file_hash(item)
            if file_hash:
                if file_hash not in hash_map:
                    hash_map[file_hash] = []
                hash_map[file_hash].append(item)

    # Find duplicates (keep first, mark rest as duplicates)
    duplicates_to_delete = []
    total_size = 0
    log_entries = []

    for file_hash, files in hash_map.items():
        if len(files) > 1:
            # Sort by modification time, keep oldest (original)
            files.sort(key=lambda f: f.stat().st_mtime)
            for dup in files[1:]:  # Skip the first (original)
                size = dup.stat().st_size
                duplicates_to_delete.append({
                    "file": dup.name,
                    "path": str(dup),
                    "size_mb": size / (1024 * 1024),
                    "original": files[0].name
                })
                total_size += size

                if not preview:
                    # Move to Trash instead of permanent delete
                    logger.info(f"TRASH DUPLICATE: {dup.name} (original: {files[0].name})")
                    trash_dest = move_to_trash(dup)
                    log_entries.append(f"{datetime.now().isoformat()} | TRASHED | {dup} | Original: {files[0].name}")

    total_size_mb = total_size / (1024 * 1024)

    # Write log file if we moved files
    if not preview and log_entries:
        log_file = Path(__file__).parent / "duplicates_log.txt"
        with open(log_file, "a") as f:
            f.write(f"\n--- Session: {datetime.now().isoformat()} ---\n")
            f.write("\n".join(log_entries) + "\n")

    if preview:
        msg = f"Preview: {len(duplicates_to_delete)} duplicates found ({total_size_mb:.1f} MB)"
    else:
        msg = f"Moved {len(duplicates_to_delete)} duplicates to Trash ({total_size_mb:.1f} MB). Check Trash to restore if needed."

    return DuplicatesResponse(
        message=msg,
        duplicates=duplicates_to_delete,
        total_size_mb=total_size_mb,
        preview_mode=preview
    )

# Backwards compatible wrapper
def find_duplicates(preview=True):
    return find_duplicates_in(DOWNLOADS_PATH, preview)

def organize_downloads(preview=True):
    return organize_folder(DOWNLOADS_PATH, preview)

@app.get("/api/duplicates/preview")
async def preview_duplicates():
    """Preview duplicate files (no action taken)"""
    return find_duplicates(preview=True)

@app.post("/api/duplicates/delete")
async def delete_duplicates():
    """Delete duplicate files (requires explicit call)"""
    return find_duplicates(preview=False)

# ============ SPORTS CARD TRACKER ============

class Card(BaseModel):
    name: str
    year: Optional[str] = None
    brand: Optional[str] = None  # Topps, Panini, Upper Deck, etc.
    sport: Optional[str] = None  # Baseball, Basketball, Football, Hockey
    player: Optional[str] = None
    condition: Optional[str] = None  # Raw, PSA 10, BGS 9.5, etc.
    purchase_price: Optional[float] = None
    estimated_value: Optional[float] = None
    notes: Optional[str] = None
    image_path: Optional[str] = None
    added_date: Optional[str] = None
    last_price_check: Optional[str] = None

class CardResponse(BaseModel):
    id: int
    card: dict
    message: str

@app.get("/api/cards")
async def list_cards():
    """List all cards in collection"""
    memory = load_memory()
    cards = memory.get("cards", [])
    total_value = sum(c.get("estimated_value", 0) or 0 for c in cards)
    total_cost = sum(c.get("purchase_price", 0) or 0 for c in cards)
    return {
        "cards": cards,
        "total_cards": len(cards),
        "total_estimated_value": round(total_value, 2),
        "total_cost": round(total_cost, 2),
        "profit_loss": round(total_value - total_cost, 2)
    }

@app.post("/api/cards")
async def add_card(card: Card):
    """Add a new card to collection"""
    memory = load_memory()
    card_dict = card.dict()
    card_dict["id"] = len(memory["cards"]) + 1
    card_dict["added_date"] = datetime.now().isoformat()
    memory["cards"].append(card_dict)
    save_memory(memory)
    logger.info(f"CARD ADDED: {card.name}")
    return {"message": f"Added {card.name} to collection", "card": card_dict}

@app.delete("/api/cards/{card_id}")
async def delete_card(card_id: int):
    """Remove a card from collection"""
    memory = load_memory()
    cards = memory.get("cards", [])
    card_to_delete = None
    for i, c in enumerate(cards):
        if c.get("id") == card_id:
            card_to_delete = cards.pop(i)
            break
    if card_to_delete:
        save_memory(memory)
        logger.info(f"CARD DELETED: {card_to_delete.get('name')}")
        return {"message": f"Removed {card_to_delete.get('name')} from collection"}
    raise HTTPException(status_code=404, detail="Card not found")

@app.post("/api/cards/identify")
async def identify_card(file: UploadFile = File(...)):
    """Use Claude vision to identify a sports card from photo"""
    file_bytes = await file.read()

    # Compress image if needed
    from io import BytesIO
    try:
        from PIL import Image
        img = Image.open(BytesIO(file_bytes))
        if img.mode in ('RGBA', 'P'):
            img = img.convert('RGB')

        max_size = 4 * 1024 * 1024
        quality = 85
        output = BytesIO()

        while True:
            output.seek(0)
            output.truncate()
            img.save(output, format='JPEG', quality=quality)
            if output.tell() <= max_size or quality <= 20:
                break
            if quality > 40:
                quality -= 15
            else:
                img = img.resize((int(img.width * 0.7), int(img.height * 0.7)), Image.LANCZOS)
                quality = 70

        file_bytes = output.getvalue()
    except ImportError:
        pass

    base64_image = base64.standard_b64encode(file_bytes).decode('utf-8')

    # Ask Claude to identify the card
    try:
        response = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1024,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": "image/jpeg",
                            "data": base64_image
                        }
                    },
                    {
                        "type": "text",
                        "text": """Identify this sports card. Provide the following in JSON format:
{
    "player": "Player name",
    "year": "Year of the card",
    "brand": "Card brand (Topps, Panini, Upper Deck, etc.)",
    "set_name": "Specific set name if identifiable",
    "card_number": "Card number if visible",
    "sport": "Sport (Baseball, Basketball, Football, Hockey, etc.)",
    "rookie": true/false if this is a rookie card,
    "parallel": "Any parallel/variant info (refractor, prizm, numbered, etc.)",
    "condition_notes": "Any visible condition issues",
    "confidence": "high/medium/low"
}

If you cannot identify some fields, set them to null. Be accurate - this is for collection tracking."""
                    }
                ]
            }]
        )

        # Parse response
        response_text = response.content[0].text
        # Extract JSON from response
        import re
        json_match = re.search(r'\{[^{}]*\}', response_text, re.DOTALL)
        if json_match:
            card_info = json.loads(json_match.group())
            logger.info(f"CARD IDENTIFIED: {card_info.get('player', 'Unknown')} {card_info.get('year', '')}")
            return {"identified": True, "card_info": card_info, "raw_response": response_text}
        else:
            return {"identified": False, "raw_response": response_text}

    except Exception as e:
        logger.error(f"Card identification error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

async def scrape_card_prices(search_query: str):
    """Scrape real prices from eBay with clickable links"""
    from playwright.async_api import async_playwright

    sales = []

    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page(
                user_agent="Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"
            )

            # Search eBay sold listings
            ebay_url = f"https://www.ebay.com/sch/i.html?_nkw={search_query.replace(' ', '+')}&LH_Complete=1&LH_Sold=1&_sop=13"
            logger.info(f"Scraping: {ebay_url}")

            await page.goto(ebay_url, timeout=30000)
            await page.wait_for_timeout(3000)

            # Get full page HTML and parse with regex
            html = await page.content()

            # Find all item blocks - look for item URLs and nearby prices
            # Pattern: find /itm/NUMBER and the price near it
            item_pattern = r'href="(https://www\.ebay\.com/itm/(\d+))[^"]*"[^>]*>([^<]+)</a>.*?\$(\d+\.?\d*)'

            # Simpler approach: find all item IDs and prices separately
            item_ids = re.findall(r'ebay\.com/itm/(\d{12,})', html)
            prices = re.findall(r'\$(\d{1,3}(?:,\d{3})*\.\d{2})', html)

            # Get titles - look for listing titles
            titles = re.findall(r'class="s-item__title"[^>]*>(?:<span[^>]*>)?([^<]+)', html)

            logger.info(f"Found {len(item_ids)} items, {len(prices)} prices, {len(titles)} titles")

            # Match them up (they should be in order on the page)
            seen_ids = set()
            price_idx = 0
            title_idx = 0

            for item_id in item_ids[:25]:
                if item_id in seen_ids:
                    continue
                seen_ids.add(item_id)

                # Skip if not enough data
                if price_idx >= len(prices):
                    break

                # Get price
                try:
                    price = float(prices[price_idx].replace(',', ''))
                    price_idx += 1
                except:
                    continue

                # Skip unreasonable prices
                if price < 1 or price > 50000:
                    continue

                # Get title
                title = "eBay listing"
                if title_idx < len(titles):
                    t = titles[title_idx].strip()
                    if t and "Shop on eBay" not in t and len(t) > 5:
                        title = t[:60]
                    title_idx += 1

                url = f"https://www.ebay.com/itm/{item_id}"

                sales.append({
                    "price": price,
                    "title": title,
                    "url": url,
                    "source": "eBay"
                })

            await browser.close()
            logger.info(f"eBay: Extracted {len(sales)} sales")

    except Exception as e:
        logger.error(f"Scrape error: {e}")

    # Remove outliers if we have enough data
    if len(sales) >= 5:
        prices = sorted([s["price"] for s in sales])
        q1, q3 = prices[len(prices)//4], prices[3*len(prices)//4]
        iqr = q3 - q1
        sales = [s for s in sales if (q1 - 1.5*iqr) <= s["price"] <= (q3 + 1.5*iqr)]

    return sales

async def get_card_prices(card_info: dict):
    """Get real prices from web scraping with individual listing links"""

    # Build search query
    search_parts = []
    if card_info.get("year"):
        search_parts.append(card_info["year"])
    if card_info.get("brand"):
        search_parts.append(card_info["brand"])
    if card_info.get("player"):
        search_parts.append(card_info["player"])
    if card_info.get("condition"):
        search_parts.append(card_info["condition"])

    search_query = " ".join(search_parts)

    # Scrape real prices - now returns list of sales with URLs
    sales = await scrape_card_prices(search_query)

    result = {
        "search_query": search_query,
        "sales_found": len(sales),
        "listings": []
    }

    if sales:
        prices = sorted([s["price"] for s in sales])
        result["low"] = round(min(prices), 2)
        result["high"] = round(max(prices), 2)
        result["median"] = round(prices[len(prices) // 2], 2)
        result["average"] = round(sum(prices) / len(prices), 2)

        # Include individual listings for verification
        result["listings"] = [
            {"price": s["price"], "title": s["title"], "url": s["url"]}
            for s in sales[:15]
        ]

    return result

@app.get("/api/cards/{card_id}/price")
async def lookup_card_price(card_id: int):
    """Get real prices for a card by scraping eBay sold listings with verification links"""
    memory = load_memory()
    cards = memory.get("cards", [])
    card = None
    card_index = None
    for i, c in enumerate(cards):
        if c.get("id") == card_id:
            card = c
            card_index = i
            break

    if not card:
        raise HTTPException(status_code=404, detail="Card not found")

    # Get real prices from web scraping
    price_data = await get_card_prices(card)

    search_query = price_data.get("search_query", "")

    result = {
        "card_id": card_id,
        "card_name": card.get("name"),
        "last_price_check": datetime.now().isoformat(),
        "search_query": search_query,
        "sales_found": price_data.get("sales_found", 0),
        "ebay_search": f"https://www.ebay.com/sch/i.html?_nkw={search_query.replace(' ', '+')}&LH_Complete=1&LH_Sold=1"
    }

    if price_data.get("sales_found", 0) > 0:
        result["low"] = price_data.get("low")
        result["high"] = price_data.get("high")
        result["median"] = price_data.get("median")
        result["average"] = price_data.get("average")
        result["listings"] = price_data.get("listings", [])  # Individual listings with URLs

        # Update card's estimated value with median
        if card_index is not None:
            memory["cards"][card_index]["estimated_value"] = price_data.get("median")
            memory["cards"][card_index]["last_price_check"] = result["last_price_check"]
            save_memory(memory)
            result["value_updated"] = True

        logger.info(f"PRICE CHECK: {card.get('name')} - ${price_data.get('median')} median ({price_data.get('sales_found')} sales)")
    else:
        result["message"] = "No sales found. Check the eBay link manually."

    return result

@app.post("/api/cards/price-check-all")
async def check_all_card_prices():
    """Check prices for all cards with verification links"""
    memory = load_memory()
    cards = memory.get("cards", [])

    if not cards:
        return {"message": "No cards in collection"}

    results = []
    for card in cards:
        price_data = await get_card_prices(card)

        if price_data.get("sales_found", 0) > 0:
            old_value = card.get("estimated_value", 0) or 0
            new_value = price_data.get("median")
            card["estimated_value"] = new_value
            card["last_price_check"] = datetime.now().isoformat()

            change = new_value - old_value
            results.append({
                "card": card.get("name"),
                "old_value": old_value,
                "new_value": new_value,
                "change": round(change, 2),
                "sales_found": price_data.get("sales_found"),
                "listings": price_data.get("listings", [])[:5]  # Top 5 for verification
            })
        else:
            results.append({
                "card": card.get("name"),
                "error": "No sales found"
            })

    save_memory(memory)

    total_value = sum(c.get("estimated_value", 0) or 0 for c in cards)
    total_cost = sum(c.get("purchase_price", 0) or 0 for c in cards)

    logger.info(f"BULK PRICE CHECK: {len(cards)} cards, total value ${total_value:.2f}")

    return {
        "cards_checked": len(cards),
        "results": results,
        "total_value": round(total_value, 2),
        "total_cost": round(total_cost, 2),
        "profit_loss": round(total_value - total_cost, 2)
    }

# ============ AI LIBRARY ============

@app.get("/api/library")
async def list_library():
    """List all documents in AI library"""
    memory = load_memory()
    library = memory.get("ai_library", [])

    # Also list actual files in library folder
    files = []
    if LIBRARY_PATH.exists():
        for f in LIBRARY_PATH.iterdir():
            if f.is_file() and not f.name.startswith('.'):
                files.append({
                    "filename": f.name,
                    "size_kb": round(f.stat().st_size / 1024, 1),
                    "modified": datetime.fromtimestamp(f.stat().st_mtime).isoformat()
                })

    return {
        "documents": library,
        "files": files,
        "total": len(library) + len(files)
    }

@app.post("/api/library/upload")
async def upload_to_library(
    file: UploadFile = File(...),
    title: str = Form(default="")
):
    """Upload a document to the AI library"""
    file_bytes = await file.read()

    # Save file to library folder
    safe_name = re.sub(r'[^\w\-_\.]', '_', file.filename)
    save_path = LIBRARY_PATH / safe_name
    save_path.write_bytes(file_bytes)

    # Add to memory index
    memory = load_memory()
    doc_title = title if title else file.filename
    memory["ai_library"].append({
        "title": doc_title,
        "filename": safe_name,
        "added": datetime.now().isoformat(),
        "size_kb": round(len(file_bytes) / 1024, 1)
    })
    save_memory(memory)

    logger.info(f"LIBRARY UPLOAD: {doc_title}")

    return {"message": f"Added '{doc_title}' to library", "filename": safe_name}

@app.get("/api/library/{filename}")
async def get_library_doc(filename: str):
    """Get a document from the library"""
    file_path = LIBRARY_PATH / filename
    if file_path.exists():
        return FileResponse(file_path)
    raise HTTPException(status_code=404, detail="Document not found")

@app.delete("/api/library/{filename}")
async def delete_library_doc(filename: str):
    """Remove a document from the library"""
    file_path = LIBRARY_PATH / filename
    if file_path.exists():
        file_path.unlink()

    # Remove from memory index
    memory = load_memory()
    memory["ai_library"] = [d for d in memory["ai_library"] if d.get("filename") != filename]
    save_memory(memory)

    logger.info(f"LIBRARY DELETE: {filename}")
    return {"message": f"Removed '{filename}' from library"}

@app.get("/api/ai/progress")
async def get_ai_progress():
    """Get AI learning progress"""
    memory = load_memory()
    learning = memory.get("ai_learning", {})
    library = memory.get("ai_library", [])

    return {
        "skills_learned": learning.get("skills", []),
        "ideas_saved": learning.get("ideas", []),
        "topics_explored": learning.get("topics_explored", []),
        "library_docs": len(library),
        "total_skills": len(learning.get("skills", []))
    }

@app.delete("/api/ai/progress")
async def reset_ai_progress():
    """Reset AI learning progress"""
    memory = load_memory()
    memory["ai_learning"] = {
        "topics_explored": [],
        "skills": [],
        "ideas": [],
        "resources": []
    }
    save_memory(memory)
    return {"message": "AI learning progress reset"}

@app.get("/api/cards/search")
async def search_card_price(q: str):
    """Quick price search without adding card to collection"""
    if not q or len(q) < 3:
        raise HTTPException(status_code=400, detail="Search query too short")

    logger.info(f"QUICK SEARCH: {q}")

    # Scrape prices
    prices, sources = await scrape_card_prices(q)

    result = {
        "search_query": q,
        "prices_found": len(prices),
        "sources": list(set(sources)),
        "ebay_link": f"https://www.ebay.com/sch/i.html?_nkw={q.replace(' ', '+')}&LH_Complete=1&LH_Sold=1"
    }

    if prices:
        prices.sort()
        result["low"] = round(min(prices), 2)
        result["high"] = round(max(prices), 2)
        result["median"] = round(prices[len(prices) // 2], 2)
        result["average"] = round(sum(prices) / len(prices), 2)
        result["recent_sales"] = [round(p, 2) for p in prices[:10]]

        logger.info(f"QUICK SEARCH RESULT: {q} - ${result['median']} median")
    else:
        result["message"] = "No prices found. Check eBay link manually."

    return result

@app.get("/api/health")
async def health():
    return {"status": "ok", "service": "HazardIQ", "phase": 1}

# ============ OHS STUDY PLATFORM ============

from anthropic import AsyncAnthropic
from fastapi.responses import StreamingResponse
import asyncio

async_client = AsyncAnthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

OHS_DATA_DIR = DATA_DIR / "ohs_data"
OHS_DATA_DIR.mkdir(exist_ok=True)
OHS_UPLOADS_DIR = OHS_DATA_DIR / "uploads"
OHS_UPLOADS_DIR.mkdir(exist_ok=True)

def ohs_load(file, fallback):
    p = OHS_DATA_DIR / file
    if not p.exists():
        return fallback() if callable(fallback) else json.loads(json.dumps(fallback))
    try:
        return json.loads(p.read_text())
    except:
        return fallback() if callable(fallback) else json.loads(json.dumps(fallback))

def ohs_save(file, data):
    (OHS_DATA_DIR / file).write_text(json.dumps(data, indent=2))

OHS_CHAPTERS = [
    {"id": 100, "code": "OHS 100", "title": "Introduction to OHS", "description": "Internal Responsibility System, worker and employer duties, right to refuse, JHSC, and Canadian OHS legislation overview.", "color": "#94a3b8"},
    {"id": 105, "code": "OHS 105", "title": "Organizational Development & Behaviour", "description": "Organizational theory, workplace culture, team dynamics, motivation, change management, and OHS integration.", "color": "#fb7185"},
    {"id": 110, "code": "OHS 110", "title": "Leadership and Communications", "description": "Safety leadership, communication strategies, worker engagement, behavior-based safety, and change management.", "color": "#a78bfa"},
    {"id": 115, "code": "OHS 115", "title": "Incident Investigation & Response", "description": "Root cause analysis, investigation methods, ICAM, reporting requirements, corrective actions, and prevention.", "color": "#f97316"},
    {"id": 120, "code": "OHS 120", "title": "Safety Training", "description": "OHS foundations, hazard types, workplace inspections, safety culture, and training program design.", "color": "#4ade80"},
    {"id": 125, "code": "OHS 125", "title": "Safety Management Systems", "description": "SMS frameworks, ISO 45001, PDCA cycle, auditing, continuous improvement, and program management.", "color": "#34d399"},
    {"id": 130, "code": "OHS 130", "title": "Hazard ID, Risk Assessment & Controls", "description": "WHMIS 2015, GHS, SDS, TDG, hierarchy of controls, risk matrices, and spill response.", "color": "#f59e0b"},
    {"id": 135, "code": "OHS 135", "title": "Law & Ethics", "description": "Canadian OHS legislation, regulatory framework, ethical decision-making, professional liability, and due diligence.", "color": "#e879f9"},
    {"id": 140, "code": "OHS 140", "title": "Industrial/Occupational Hygiene", "description": "Recognition, evaluation, and control of chemical, physical, and biological workplace hazards.", "color": "#60a5fa"},
    {"id": 145, "code": "OHS 145", "title": "Fire Management", "description": "Fire behavior, prevention, protection systems, emergency response, and Canadian OHS law.", "color": "#ff5c26"},
]

OHS_DEFAULT_TOPICS = {
    100: [{"id":1,"title":"Introduction to OHS in Canada","subtitle":"History, scope & why it matters"},{"id":2,"title":"Internal Responsibility System","subtitle":"Foundation of Canadian OHS"},{"id":3,"title":"Worker Rights & Duties","subtitle":"Right to know, participate & refuse"},{"id":4,"title":"Employer & Supervisor Duties","subtitle":"Due diligence & duty of care"},{"id":5,"title":"Joint Health & Safety Committees","subtitle":"Structure, powers & effectiveness"},{"id":6,"title":"Canadian OHS Legislation Overview","subtitle":"Federal, provincial & territorial"}],
    105: [{"id":1,"title":"Organizational Theory & Structure","subtitle":"How organizations work & why it matters for OHS"},{"id":2,"title":"Workplace Culture & Climate","subtitle":"Safety culture, values & norms"},{"id":3,"title":"Team Dynamics & Group Behaviour","subtitle":"Group decision-making, conformity & conflict"},{"id":4,"title":"Motivation & Human Behaviour","subtitle":"Theories of motivation & behaviour change"},{"id":5,"title":"Change Management","subtitle":"Managing organizational change & resistance"},{"id":6,"title":"OHS Integration in Organizations","subtitle":"Building OHS into organizational systems"}],
    110: [{"id":1,"title":"Safety Leadership Principles","subtitle":"Traits, styles & transformational leadership"},{"id":2,"title":"Communication in OHS","subtitle":"Effective messaging, barriers & strategies"},{"id":3,"title":"Worker Engagement & Participation","subtitle":"Consultation, involvement & empowerment"},{"id":4,"title":"Behavior-Based Safety","subtitle":"Observation, feedback & behaviour change"},{"id":5,"title":"Conflict & Difficult Conversations","subtitle":"Navigating disagreements & safety non-compliance"},{"id":6,"title":"Leading Change in Safety Culture","subtitle":"Change management & continuous improvement"}],
    115: [{"id":1,"title":"Incident Investigation Principles","subtitle":"Purpose, scope & legal requirements"},{"id":2,"title":"Root Cause Analysis Methods","subtitle":"5 Whys, fishbone, fault tree & ICAM"},{"id":3,"title":"Investigation Process & Procedures","subtitle":"Scene preservation, interviews & evidence"},{"id":4,"title":"Reporting & Documentation","subtitle":"What, when & how to report"},{"id":5,"title":"Corrective & Preventive Actions","subtitle":"From findings to lasting fixes"},{"id":6,"title":"Near Misses & Hazard Reporting","subtitle":"Building a reporting culture"}],
    120: [{"id":1,"title":"OHS Foundations & Hazard Types","subtitle":"Physical, chemical, biological, ergonomic & psychosocial"},{"id":2,"title":"Workplace Inspections","subtitle":"Planning, conducting & reporting inspections"},{"id":3,"title":"Safety Culture & Program Development","subtitle":"Building & sustaining safety programs"},{"id":4,"title":"Training Program Design","subtitle":"Needs assessment, adult learning & evaluation"},{"id":5,"title":"Delivering Safety Training","subtitle":"Facilitation methods & engagement techniques"},{"id":6,"title":"Evaluating Training Effectiveness","subtitle":"Kirkpatrick model & continuous improvement"}],
    125: [{"id":1,"title":"SMS Frameworks & Standards","subtitle":"OHSAS 18001, ISO 45001 & Canadian standards"},{"id":2,"title":"PDCA Cycle in Safety Management","subtitle":"Plan, Do, Check, Act for continuous improvement"},{"id":3,"title":"Hazard & Risk Management","subtitle":"Systematic identification, assessment & control"},{"id":4,"title":"Safety Program Elements","subtitle":"Policies, procedures, roles & responsibilities"},{"id":5,"title":"Auditing & Performance Measurement","subtitle":"Internal audits, KPIs & lagging/leading indicators"},{"id":6,"title":"Continuous Improvement","subtitle":"CAPA systems, lessons learned & benchmarking"}],
    130: [{"id":1,"title":"WHMIS 2015 & GHS","subtitle":"Classification, labels & Hazardous Products Act"},{"id":2,"title":"Safety Data Sheets (SDS)","subtitle":"Reading & interpreting the 16 sections"},{"id":3,"title":"Hazard Identification Methods","subtitle":"JHA, HAZOP, what-if & checklists"},{"id":4,"title":"Risk Assessment & Risk Matrices","subtitle":"Likelihood, severity & risk ranking"},{"id":5,"title":"Hierarchy of Controls","subtitle":"Elimination through to PPE"},{"id":6,"title":"Transportation of Dangerous Goods","subtitle":"TDG Act, classes, placards & docs"},{"id":7,"title":"Spill Response & Emergency Procedures","subtitle":"Containment, cleanup & reporting"}],
    135: [{"id":1,"title":"Canadian OHS Legal Framework","subtitle":"Federal, provincial & territorial jurisdiction"},{"id":2,"title":"OHS Legislation & Regulations","subtitle":"Key acts, regulations & standards"},{"id":3,"title":"Enforcement & Compliance","subtitle":"Inspections, orders, penalties & Bill C-45"},{"id":4,"title":"Ethics in OHS Practice","subtitle":"Ethical frameworks & professional obligations"},{"id":5,"title":"Professional Liability & Due Diligence","subtitle":"Demonstrating due diligence & avoiding liability"},{"id":6,"title":"Worker Rights & Compensation","subtitle":"WCB/WSIB, return to work & appeals"}],
    140: [{"id":1,"title":"Introduction to Occupational Hygiene","subtitle":"Anticipation, recognition, evaluation, control"},{"id":2,"title":"Chemical Hazards","subtitle":"Gases, vapours, dusts, fumes & mists"},{"id":3,"title":"Physical Hazards — Noise & Vibration","subtitle":"Measurement, limits & hearing conservation"},{"id":4,"title":"Physical Hazards — Radiation & Thermal","subtitle":"Ionizing, non-ionizing, heat & cold stress"},{"id":5,"title":"Biological Hazards","subtitle":"Bloodborne pathogens, mould, infectious agents"},{"id":6,"title":"Exposure Assessment & Monitoring","subtitle":"TLVs, OELs, sampling & instruments"},{"id":7,"title":"Engineering & Administrative Controls","subtitle":"Ventilation, substitution, work practices"},{"id":8,"title":"PPE & Hygiene Programs","subtitle":"Respiratory protection, program management"}],
    145: [{"id":1,"title":"Fire Behavior","subtitle":"How fires start, spread & kill"},{"id":2,"title":"Fire Prevention & Hazard ID","subtitle":"Stop fires before they start"},{"id":3,"title":"Fire Protection Systems","subtitle":"Sprinklers, alarms & extinguishers"},{"id":4,"title":"Emergency Response Procedures","subtitle":"When the alarm goes off"},{"id":5,"title":"Human Behavior in Fire","subtitle":"Why people die — and how to prevent it"},{"id":6,"title":"OHS Legal Responsibilities","subtitle":"Canadian law & your duty of care"},{"id":7,"title":"Fire Risk Assessment","subtitle":"Identify, evaluate, control"},{"id":8,"title":"Digital EHS Tools","subtitle":"Technology for modern safety officers"}],
}

OHS_CHAPTER_CONTEXT = {
    145: "\nOHS 145: Fire Management. Fire behavior, prevention, hazard ID, emergency response, protection systems, human behavior in fire, legal responsibilities under Canadian OHS law.",
    140: "\nOHS 140: Industrial/Occupational Hygiene. Chemical, physical, biological hazard recognition, evaluation and control. TLVs/OELs, sampling, ventilation, PPE.",
    135: "\nOHS 135: Law & Ethics. Canadian OHS legislation, regulatory framework, ethical decision-making, professional liability, due diligence, worker rights.",
    130: "\nOHS 130: Hazard Identification, Risk Assessment & Controls. WHMIS 2015, GHS, SDS, TDG, hierarchy of controls, risk matrices, spill response.",
    125: "\nOHS 125: Safety Management Systems. SMS frameworks, OHSAS 18001/ISO 45001, PDCA cycle, auditing, continuous improvement, program management.",
    120: "\nOHS 120: Safety Training — Introduction to the Fundamental Principles. OHS foundations, hazard types, workplace inspections, safety culture, training program design.",
    115: "\nOHS 115: Incident Investigation & Response. Root cause analysis, investigation methods, ICAM, reporting requirements, corrective actions, prevention.",
    110: "\nOHS 110: Leadership and Communications. Safety leadership, communication strategies, worker engagement, behavior-based safety, change management.",
    105: "\nOHS 105: Organizational Development & Behaviour. Organizational theory, workplace culture, team dynamics, motivation, change management, OHS integration.",
    100: "\nOHS 100: Introduction to OHS. Internal Responsibility System, worker and employer duties, right to refuse, JHSC, OHS legislation overview, Canadian context.",
}

OHS_BASE_PROMPT = """You are a world-leading professor in OHS, built into HazardIQ — a Canadian-focused OHS study system.

Transform course material into engaging, practical study content. Use these EXACT section headers in order:
## KEY CONCEPTS
## WHY IT MATTERS
## COMMON MISTAKES STUDENTS MAKE
## THINK LIKE A SAFETY OFFICER
## MEMORY TOOLS
## RAPID REVIEW

Include 2 Canadian workplace scenarios, mnemonics, bullet-point rapid review. Reference National Fire Code of Canada, provincial OHS legislation, WorkSafeBC. Label difficulty: [BASIC], [INTERMEDIATE], [CHALLENGE]. Tone: Direct, practical, no fluff.

FIRE TOPICS RULE: Whenever the topic involves fire classification, fire suppression, extinguishing agents, or fire prevention:
1. When presenting the fire classification table (Class A/B/C/D/K), ALWAYS use exactly these 5 columns:
   | Class | Fuel Type | Canadian Workplace Example | Extinguishing Agent(s) | How It Works |
   Fill the "How It Works" column with a precise one-sentence mechanism for each class:
   - Class A: Water removes heat; ABC dry chemical interrupts the free-radical chain reaction
   - Class B: CO₂ displaces oxygen (smothers); ABC/BC dry chemical interrupts chain reaction; foam blankets surface suppressing vapours
   - Class C: CO₂ or clean agent (FM-200/Novec) — non-conductive, no residue; NEVER water (electrocution risk)
   - Class D: Class-D-specific dry powder smothers and absorbs heat — NEVER water, CO₂, or standard dry chemical (violent reactions)
   - Class K: Wet chemical causes saponification (soap reaction with hot oil), sealing the surface and cooling below auto-ignition
2. After the table include a subsection "HOW EXTINGUISHING AGENTS WORK" covering each agent's mechanism and which leg of the fire triangle/tetrahedron it attacks.
3. Always note what NOT to use on each class and why.

CRITICAL — HYPERLINKS: Whenever you mention a specific piece of legislation, regulation, standard, government body, or official document, format it as a markdown hyperlink to the real official website. Only link to real, accurate URLs."""

DISTRACTOR_RULE = "DISTRACTOR QUALITY RULE: all 4 options plausible and similar length. Never make correct answer longest. No obviously wrong distractors."

MAX_SOURCE_CHARS = 80000
CHARS_PER_SOURCE = 12000

def ohs_score_relevance(text, query):
    if not query:
        return 0
    words = [w for w in query.lower().split() if len(w) > 3]
    t = text.lower()
    return sum(t.count(w) for w in words)

def ohs_get_source_context(cid, query=""):
    sources = (ohs_load("sources.json", {}) or {}).get(str(cid), [])
    if not sources:
        return ""
    if query and len(sources) > 5:
        scored = sorted(sources, key=lambda x: ohs_score_relevance(x.get("text",""), query), reverse=True)
        selected, budget = [], MAX_SOURCE_CHARS
        for s in scored:
            if budget <= 0:
                break
            chunk = s["text"][:min(CHARS_PER_SOURCE, budget)]
            selected.append((s["name"], chunk))
            budget -= len(chunk)
    else:
        per = min(CHARS_PER_SOURCE, MAX_SOURCE_CHARS // max(len(sources),1))
        selected = [(s["name"], s["text"][:per]) for s in sources]
    return "\n\n─── UPLOADED MATERIAL ───\n" + "\n\n".join(f"--- {n} ---\n{t}" for n,t in selected)

def ohs_build_system(cid, query=""):
    return OHS_BASE_PROMPT + OHS_CHAPTER_CONTEXT.get(int(cid) if str(cid).isdigit() else 0, "") + ohs_get_source_context(cid, query)

async def ohs_extract_text(file_bytes: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext in ['.txt', '.md', '.csv', '.tsv', '.log', '.json']:
        return file_bytes.decode('utf-8', errors='replace')
    if ext == '.pdf':
        import fitz
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        text = "".join(page.get_text() for page in pdf)
        pdf.close()
        return text
    if ext == '.docx':
        from docx import Document
        import io
        doc = Document(io.BytesIO(file_bytes))
        return '\n'.join(p.text for p in doc.paragraphs)
    return file_bytes.decode('utf-8', errors='replace')

async def ohs_sse_stream(model, max_tokens, system, messages):
    async def gen():
        try:
            async with async_client.messages.stream(model=model, max_tokens=max_tokens, system=system, messages=messages) as stream:
                async for text in stream.text_stream:
                    yield f"data: {json.dumps({'text': text})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
    return StreamingResponse(gen(), media_type="text/event-stream", headers={"Cache-Control": "no-cache", "Connection": "keep-alive"})

# Seed default topics on startup
_saved_topics = ohs_load("topics.json", {})
_merged = {**{str(k): v for k,v in OHS_DEFAULT_TOPICS.items()}, **_saved_topics}
ohs_save("topics.json", _merged)

# ── Chapters & Topics ─────────────────────────────────────────────────────────
@app.get("/api/ohs/chapters")
async def ohs_chapters():
    topics = ohs_load("topics.json", {})
    return [{**c, "topicCount": len(topics.get(str(c["id"]), []))} for c in OHS_CHAPTERS]

@app.get("/api/ohs/chapters/{cid}/topics")
async def ohs_get_topics(cid: int):
    return ohs_load("topics.json", {}).get(str(cid), [])

@app.post("/api/ohs/chapters/{cid}/topics")
async def ohs_add_topic(cid: int, body: dict):
    title = body.get("title","").strip()
    if not title:
        raise HTTPException(400, "title required")
    t = ohs_load("topics.json", {})
    key = str(cid)
    if key not in t:
        t[key] = []
    nid = max((x["id"] for x in t[key]), default=0) + 1
    topic = {"id": nid, "title": title, "subtitle": body.get("subtitle","")}
    t[key].append(topic)
    ohs_save("topics.json", t)
    return topic

@app.delete("/api/ohs/chapters/{cid}/topics/{tid}")
async def ohs_del_topic(cid: int, tid: int):
    t = ohs_load("topics.json", {})
    key = str(cid)
    if key in t:
        t[key] = [x for x in t[key] if x["id"] != tid]
        ohs_save("topics.json", t)
    return {"ok": True}

# ── Sources ───────────────────────────────────────────────────────────────────
@app.get("/api/ohs/sources/{cid}")
async def ohs_list_sources(cid: str):
    s = ohs_load("sources.json", {}).get(cid, [])
    return [{"id":x["id"],"name":x["name"],"type":x["type"],"addedAt":x["addedAt"],"textLength":len(x["text"])} for x in s]

@app.post("/api/ohs/sources/{cid}/upload")
async def ohs_upload_source(cid: str, file: UploadFile = File(...)):
    file_bytes = await file.read()
    try:
        text = await ohs_extract_text(file_bytes, file.filename)
        if not text.strip():
            raise ValueError("No text extracted")
    except Exception as e:
        raise HTTPException(400, str(e))
    s = ohs_load("sources.json", {})
    if cid not in s:
        s[cid] = []
    xid = f"{int(datetime.now().timestamp()*1000):x}{len(s[cid]):x}"
    entry = {"id": xid, "name": file.filename, "type": "file", "text": text[:100000], "addedAt": datetime.now().isoformat()}
    s[cid].append(entry)
    ohs_save("sources.json", s)
    return {"id": xid, "name": file.filename, "type": "file", "addedAt": entry["addedAt"], "textLength": len(entry["text"])}

@app.delete("/api/ohs/sources/{cid}/{sid}")
async def ohs_del_source(cid: str, sid: str):
    s = ohs_load("sources.json", {})
    if cid in s:
        s[cid] = [x for x in s[cid] if x["id"] != sid]
        ohs_save("sources.json", s)
    return {"ok": True}

# ── Content generation (SSE streaming) ───────────────────────────────────────
@app.post("/api/ohs/generate")
async def ohs_generate(body: dict):
    cid = body.get("chapterId")
    title = body.get("topicTitle","")
    tid = body.get("topicId","")
    return await ohs_sse_stream(
        "claude-sonnet-4-6", 8000,
        ohs_build_system(cid, title),
        [{"role":"user","content":f'Generate complete study module for Topic {tid}: "{title}". Include ALL sections. Canadian OHS context. Use uploaded material if available.'}]
    )

@app.post("/api/ohs/flashcards")
async def ohs_flashcards(body: dict):
    cid = body.get("chapterId")
    title = body.get("topicTitle","")
    tid = body.get("topicId","")
    try:
        msg = await async_client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=4000,
            system=f"Canadian OHS prof. Output ONLY a JSON array of flashcard objects.\n{DISTRACTOR_RULE}\n{ohs_get_source_context(cid, title)}",
            messages=[{"role":"user","content":f'20 flashcards for Topic {tid}: "{title}" (Ch {cid}). Format: [{{"id":1,"difficulty":"basic","question":"...","answer":"..."}}]. 7 basic, 8 intermediate, 5 challenge. ONLY JSON array.'}]
        )
        raw = msg.content[0].text.strip()
        m = re.search(r'\[[\s\S]*\]', raw)
        return json.loads(m.group() if m else raw)
    except Exception as e:
        raise HTTPException(500, str(e))

@app.post("/api/ohs/quiz")
async def ohs_quiz(body: dict):
    cid = body.get("chapterId")
    title = body.get("topicTitle","")
    tid = body.get("topicId","")
    try:
        msg = await async_client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=4000,
            system=f"Canadian OHS prof generating a quiz. Output ONLY a JSON array.\n{DISTRACTOR_RULE}\n{ohs_get_source_context(cid, title)}",
            messages=[{"role":"user","content":f'12 quiz questions for Topic {tid}: "{title}" (Ch {cid}). 4 multiple_choice, 4 true_false, 4 scenario. Format: {{"id","type","difficulty","question","options":[],"answer","explanation"}}. ONLY JSON array.'}]
        )
        raw = msg.content[0].text.strip()
        m = re.search(r'\[[\s\S]*\]', raw)
        return json.loads(m.group() if m else raw)
    except Exception as e:
        raise HTTPException(500, str(e))

@app.post("/api/ohs/chapters/{cid}/full-flashcards")
async def ohs_full_flashcards(cid: int):
    topics = ohs_load("topics.json", {})
    topic_list = ", ".join(t["title"] for t in topics.get(str(cid), []))
    try:
        msg = await async_client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=6000,
            system=f"Canadian OHS prof. Output ONLY a JSON array.\n{DISTRACTOR_RULE}\n{ohs_get_source_context(cid, topic_list)}",
            messages=[{"role":"user","content":f'30 flashcards covering ALL topics in Chapter {cid}: {topic_list}. Format: [{{"id":1,"difficulty":"basic","question":"...","answer":"..."}}]. 10 basic, 12 intermediate, 8 challenge. ONLY JSON array.'}]
        )
        raw = msg.content[0].text.strip()
        m = re.search(r'\[[\s\S]*\]', raw)
        return json.loads(m.group() if m else raw)
    except Exception as e:
        raise HTTPException(500, str(e))

@app.post("/api/ohs/chapters/{cid}/full-quiz")
async def ohs_full_quiz(cid: int):
    topics = ohs_load("topics.json", {})
    topic_list = ", ".join(t["title"] for t in topics.get(str(cid), []))
    try:
        msg = await async_client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=6000,
            system=f"Canadian OHS prof generating a comprehensive course exam. Output ONLY a JSON array.\n{DISTRACTOR_RULE}\n{ohs_get_source_context(cid, topic_list)}",
            messages=[{"role":"user","content":f'20 exam questions covering ALL topics in Chapter {cid}: {topic_list}. Mix of multiple_choice, true_false, scenario. Format: {{"id","type","difficulty","question","options":[],"answer","explanation"}}. ONLY JSON array.'}]
        )
        raw = msg.content[0].text.strip()
        m = re.search(r'\[[\s\S]*\]', raw)
        return json.loads(m.group() if m else raw)
    except Exception as e:
        raise HTTPException(500, str(e))

@app.post("/api/ohs/program/flashcards")
async def ohs_program_flashcards():
    courses = ", ".join(f"{c['code']}: {c['title']}" for c in OHS_CHAPTERS)
    try:
        msg = await async_client.messages.create(
            model="claude-sonnet-4-6", max_tokens=8000,
            system=f"Canadian OHS prof. Output ONLY a JSON array.\n{DISTRACTOR_RULE}",
            messages=[{"role":"user","content":f'40 flashcards covering the entire UofF OHS Safety Officer Training Program: {courses}. Format: [{{"id":1,"difficulty":"basic","question":"...","answer":"...","course":"OHS 100"}}]. 12 basic, 18 intermediate, 10 challenge. ONLY JSON array.'}]
        )
        raw = msg.content[0].text.strip()
        m = re.search(r'\[[\s\S]*\]', raw)
        return json.loads(m.group() if m else raw)
    except Exception as e:
        raise HTTPException(500, str(e))

@app.post("/api/ohs/program/quiz")
async def ohs_program_quiz():
    courses = ", ".join(f"{c['code']}: {c['title']}" for c in OHS_CHAPTERS)
    try:
        msg = await async_client.messages.create(
            model="claude-sonnet-4-6", max_tokens=8000,
            system=f"Canadian OHS prof generating a comprehensive program exam. Output ONLY a JSON array.\n{DISTRACTOR_RULE}",
            messages=[{"role":"user","content":f'30 exam questions covering the entire UofF OHS Safety Officer Training Program: {courses}. Mix of multiple_choice, true_false, scenario. Format: {{"id","type","difficulty","question","options":[],"answer","explanation","course":"OHS 100"}}. ONLY JSON array.'}]
        )
        raw = msg.content[0].text.strip()
        m = re.search(r'\[[\s\S]*\]', raw)
        return json.loads(m.group() if m else raw)
    except Exception as e:
        raise HTTPException(500, str(e))

@app.post("/api/ohs/detect-course")
async def ohs_detect_course(body: dict):
    text = body.get("text","")
    filename = body.get("filename","unknown")
    course_list = "\n".join(f"{c['id']}: {c['code']} — {c['title']}" for c in OHS_CHAPTERS)
    prompt = f'You are routing an OHS document to the right course.\nFilename: "{filename}"\nContent preview:\n{text[:2000]}\n\nCourses:\n{course_list}\n\nWhich course? Respond ONLY with JSON: {{"courseId":145,"courseName":"Fire Management","confidence":"high","reason":"one sentence"}}'
    try:
        msg = await async_client.messages.create(
            model="claude-haiku-4-5-20251001", max_tokens=150,
            messages=[{"role":"user","content":prompt}]
        )
        raw = msg.content[0].text.strip()
        m = re.search(r'\{[\s\S]*\}', raw)
        return json.loads(m.group() if m else raw)
    except Exception as e:
        raise HTTPException(500, str(e))

@app.post("/api/ohs/chat")
async def ohs_chat(body: dict):
    cid = body.get("chapterId")
    messages = body.get("messages", [])
    last_query = next((m["content"] for m in reversed(messages) if m["role"]=="user"), "")
    return await ohs_sse_stream("claude-haiku-4-5-20251001", 2000, ohs_build_system(cid, last_query), messages)

# ─── Floating Notes API ───────────────────────────────────────────────────────
NOTES_FILE = DATA_DIR / "notes.txt"

@app.get("/api/notes")
async def get_notes():
    if NOTES_FILE.exists():
        return {"content": NOTES_FILE.read_text(encoding="utf-8")}
    return {"content": ""}

class NotesBody(BaseModel):
    content: str

@app.post("/api/notes")
async def save_notes(body: NotesBody):
    NOTES_FILE.write_text(body.content, encoding="utf-8")
    return {"status": "saved"}

# Serve static files (frontend)
app.mount("/static", StaticFiles(directory=Path(__file__).parent.parent / "frontend"), name="static")
